#!/usr/bin/env python3
import os
import json
import uuid
import io
from datetime import datetime
from flask import Flask, request, jsonify, render_template_string, session, send_file, Response, stream_with_context

# Third-party conversion requirements
from pypdf import PdfReader, PdfWriter
from docx import Document
from PIL import Image

from litelaw import (
    get_system_prompt,
    call_ollama,
    call_ollama_stream,
    execute_command,
    parse_response,
    doc_write,
    MAX_CONTEXT_MESSAGES,
)

app = Flask(__name__)
app.secret_key = uuid.uuid4().hex

def _now():
    return datetime.now().strftime("%I:%M %p")

STORE_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "litelaw_store.json")

def load_store():
    """Load data from JSON file or initialize a fresh state."""
    if not os.path.exists(STORE_FILE):
        default_store = {
            "chats": {},
            "memories": [
                "User loves automation and clean terminal configurations.",
                "Always use echo for writing text to files (e.g. echo 'text' > file)."
            ],
            "documents": {
                "default-doc": {
                    "title": "Scratchpad.txt",
                    "content": "Welcome to your litelaw workspace document editor!\nYou can save temporary data, task definitions, or command templates here."
                }
            },
            "reminders": {}
        }
        save_store(default_store)
        return default_store
    try:
        with open(STORE_FILE, "r") as f:
            data = json.load(f)
            data.setdefault("reminders", {})
            return data
    except Exception:
        return {"chats": {}, "memories": [], "documents": {}, "reminders": {}}

def save_store(data):
    """Persist store dictionary cleanly back to disk."""
    with open(STORE_FILE, "w") as f:
        json.dump(data, f, indent=2)

def trim_context(session_messages):
    if len(session_messages) > MAX_CONTEXT_MESSAGES:
        system_prompt = session_messages[0]
        session_messages[:] = [system_prompt] + session_messages[-(MAX_CONTEXT_MESSAGES - 1):]

def _is_write_request(text):
    """Check if a user request involves writing text content to a file."""
    triggers = ["write", "content", "text", "echo", "save", "print", "put", "add line"]
    lower = text.lower()
    return any(t in lower for t in triggers)

_SIGNOFFS = ["have a great day", "have a nice day", "goodbye", "bye", "take care", "have fun", "see you", "thanks", "thank you"]
def _clean_answer(text):
    """Strip conversational sign-offs that the model mistakenly puts in ANSWER."""
    if not text:
        return text
    lower = text.strip().lower()
    for phrase in _SIGNOFFS:
        if lower == phrase or lower.startswith(phrase + " ") or lower.startswith(phrase + ".") or lower.startswith(phrase + "!"):
            return ""
    return text

def web_run_agent_stream(user_goal, session_messages, memories):
    """Generator that yields token + step events for real-time SSE streaming."""
    now = _now()
    session_messages.append({"role": "user", "content": f"Task: {user_goal}", "time": now})
    now = _now()

    command_history = []
    written_contents = []
    consecutive_no_action = 0
    commands_without_finish = 0
    for step in range(10):
        yield {"type": "thinking_start", "time": now}

        response_parts = []
        for token in call_ollama_stream(session_messages):
            response_parts.append(token)
            yield {"type": "token", "text": token}

        response = "".join(response_parts)
        if not response:
            yield {"type": "thinking_end"}
            yield {"type": "error", "text": "Failed to reach Ollama. Make sure it's running (ollama serve)."}
            break

        yield {"type": "thinking_end"}
        session_messages.append({"role": "assistant", "content": response, "time": now})

        thought_lines = [line.replace("THOUGHT:", "").strip() for line in response.split('\n') if line.startswith("THOUGHT:")]

        action, target = parse_response(response)
        if target:
            for typo in ['litlaw.py', 'litalaw.py', 'litalw.py', 'litalow.py', 'litelow.py', 'ltielaw.py']:
                target = target.replace(typo, 'litelaw.py')

        if action == "RUN_COMMAND" and target:
            consecutive_no_action = 0
            commands_without_finish += 1
            for typo in ['litlaw.py', 'litalaw.py', 'litalw.py', 'litalow.py', 'litelow.py', 'ltielaw.py']:
                thought_lines = [t.replace(typo, 'litelaw.py') for t in thought_lines]
            for t in thought_lines:
                yield {"type": "thought", "text": t, "time": now}
            command_history.append(target)
            repeats = sum(1 for c in command_history[-3:] if c == target)
            if repeats >= 3:
                yield {"type": "warning", "text": f"Loop detected: command '{target}' repeated {repeats} times. Forcing completion.", "time": now}
                session_messages.append({"role": "user", "content": "That command already ran successfully. The task is done. Say ACTION: FINISHED now.", "time": now})
                continue
            elif repeats >= 2:
                yield {"type": "warning", "text": f"Warning: command '{target}' is repeating. Task may already be done.", "time": now}
                yield {"type": "command", "text": target, "time": now}
                cmd_output = execute_command(target)
                yield {"type": "output", "text": cmd_output.strip(), "time": now}
                session_messages.append({"role": "user", "content": f"Command output:\n{cmd_output}\n\nNote: You already ran this command. If it worked, say ACTION: FINISHED. Do NOT repeat the same command.", "time": now})
                continue
            if commands_without_finish >= 2:
                yield {"type": "warning", "text": "Forcing completion after multiple commands.", "time": now}
                yield {"type": "command", "text": target, "time": now}
                cmd_output = execute_command(target)
                yield {"type": "output", "text": cmd_output.strip(), "time": now}
                session_messages.append({"role": "user", "content": f"Command output:\n{cmd_output}\n\nYou have run multiple commands. You MUST now summarize the results and say ACTION: FINISHED with your answer. Do NOT run more commands.", "time": now})
                continue

            if target.strip().startswith("touch ") and _is_write_request(user_goal):
                yield {"type": "warning", "text": "Touch misuse: you used 'touch' but this task requires writing content. Use echo or printf instead.", "time": now}
                session_messages.append({"role": "user", "content": "You used 'touch' which only creates empty files. This task requires writing content into the file. Use printf or echo instead. For example: printf 'content here\\n' > filename", "time": now})
                continue

            lower_target = target.lower().strip()
            if lower_target.startswith("printf ") or lower_target.startswith("echo "):
                try:
                    content = target.split(maxsplit=1)[1]
                    if ">" in content:
                        content = content.rsplit(">", 1)[0].strip()
                    content = content.strip("'\"")
                    if content and len(content) < 5000:
                        written_contents.append(content)
                except Exception:
                    pass
            if "--doc-write" in lower_target:
                try:
                    rest = target.split("--doc-write", 1)[1].strip()
                    parts = rest.split('"')
                    if len(parts) >= 3:
                        written_contents.append(parts[3])
                except Exception:
                    pass

            yield {"type": "command", "text": target, "time": now}
            cmd_output = execute_command(target)
            yield {"type": "output", "text": cmd_output.strip(), "time": now}
            session_messages.append({"role": "user", "content": f"Command output:\n{cmd_output}", "time": now})
        elif action == "FINISHED":
            consecutive_no_action = 0
            commands_without_finish = 0
            clean_target = _clean_answer(target)
            for typo in ['litlaw.py', 'litalaw.py', 'litalw.py', 'litalow.py', 'litelow.py', 'ltielaw.py']:
                clean_target = clean_target.replace(typo, 'litelaw.py')
                thought_lines = [t.replace(typo, 'litelaw.py') for t in thought_lines]
            for t in thought_lines:
                yield {"type": "thought", "text": t, "time": now}
            final_text = clean_target if clean_target else "\n".join(thought_lines)
            if final_text and written_contents:
                for wc in written_contents:
                    if wc and wc in final_text:
                        final_text = final_text.replace(wc, "").strip()
                        if not final_text:
                            final_text = "Done."
            yield {"type": "final", "text": final_text, "time": now}
            break
        else:
            consecutive_no_action += 1
            if consecutive_no_action >= 2:
                fallback = "\n".join(thought_lines) if thought_lines else re.sub(r'(ACTION|COMMAND|ANSWER):.*', '', response).strip()
                if fallback:
                    yield {"type": "final", "text": fallback, "time": now}
                else:
                    yield {"type": "warning", "text": "Formatting misalignment detected; adjusting agent constraints.", "time": now}
                break
            yield {"type": "warning", "text": "Formatting misalignment detected; adjusting agent constraints.", "time": now}
            session_messages.append({"role": "user", "content": "Invalid layout. Return your move strictly using ACTION: RUN_COMMAND or ACTION: FINISHED.", "time": now})
    else:
        yield {"type": "error", "text": "Reached step threshold limit before closure.", "time": now}

    trim_context(session_messages)


def web_run_agent(user_goal, session_messages, memories):
    """Executes the autonomous terminal workflow loop (non-streaming)."""
    return list(web_run_agent_stream(user_goal, session_messages, memories))

@app.route("/")
def index():
    return render_template_string(PAGE_HTML)

@app.route("/api/init", methods=["GET"])
def initialize_workspace():
    """Flashes out entire synchronized store data to frontend sidebar."""
    store = load_store()
    chat_list = [{"id": cid, "title": c["title"]} for cid, c in store["chats"].items()]
    doc_list = [{"id": did, "title": d["title"]} for did, d in store["documents"].items()]
    return jsonify({
        "chats": chat_list,
        "memories": store["memories"],
        "documents": doc_list
    })

@app.route("/api/chat/new", methods=["POST"])
def create_new_chat():
    store = load_store()
    chat_id = uuid.uuid4().hex
    
    # Pre-inject system prompt bundled with current persistent memories
    store["chats"][chat_id] = {
        "title": "Untitled Operations Thread",
        "messages": [{"role": "system", "content": get_system_prompt(store["memories"])}]
    }
    save_store(store)
    return jsonify({"chat_id": chat_id, "title": store["chats"][chat_id]["title"]})

@app.route("/api/chat/get", methods=["GET"])
def get_chat_history():
    chat_id = request.args.get("chat_id")
    store = load_store()
    if chat_id not in store["chats"]:
        return jsonify({"error": "Thread expired or absent"}), 404
    
    # Filter structural system records away from UI chat view nodes
    visible_steps = []
    raw_msgs = store["chats"][chat_id]["messages"]
    
    # Parse existing history back into interactive client steps
    for idx, m in enumerate(raw_msgs):
        ts = m.get("time", "")
        if m["role"] == "user" and m["content"].startswith("Task: "):
            text = m["content"].replace("Task: ", "", 1)
            if text.startswith("User instruction: "):
                text = text.replace("User instruction: ", "", 1)
            visible_steps.append({"type": "user_msg", "text": text, "time": ts})
        elif m["role"] == "assistant":
            resp = m["content"]
            thought_lines = [line.replace("THOUGHT:", "").strip() for line in resp.split('\n') if line.startswith("THOUGHT:")]
            
            # Look at subsequent indices for command execution outputs
            action, target = parse_response(resp)
            if action == "RUN_COMMAND":
                for t in thought_lines:
                    visible_steps.append({"type": "thought", "text": t, "time": ts})
                visible_steps.append({"type": "command", "text": target, "time": ts})
                if idx + 1 < len(raw_msgs) and raw_msgs[idx+1]["role"] == "user" and raw_msgs[idx+1]["content"].startswith("Command output:\n"):
                    out_text = raw_msgs[idx+1]["content"].replace("Command output:\n", "")
                    visible_steps.append({"type": "output", "text": out_text, "time": ts})
            elif action == "FINISHED":
                clean_target = _clean_answer(target)
                if thought_lines and clean_target:
                    final_text = "\n".join(thought_lines) + "\n\n" + clean_target
                elif thought_lines:
                    final_text = "\n".join(thought_lines)
                else:
                    final_text = clean_target
                visible_steps.append({"type": "final", "text": final_text, "time": ts})
                
    return jsonify({"steps": visible_steps})

@app.route("/api/chat", methods=["POST"])
def chat():
    data = request.get_json(force=True) or {}
    chat_id = data.get("chat_id")
    message = (data.get("message") or "").strip()

    if not message:
        return jsonify({"steps": [{"type": "error", "text": "Empty command syntax query submission."}]})

    store = load_store()
    
    # If client lacks an active session ID or thread got deleted, spin a new one up
    if not chat_id or chat_id not in store["chats"]:
        chat_id = uuid.uuid4().hex
        store["chats"][chat_id] = {
            "title": message[:32] + "..." if len(message) > 32 else message,
            "messages": [{"role": "system", "content": get_system_prompt(store["memories"])}]
        }
    elif store["chats"][chat_id]["title"] == "Untitled Operations Thread":
        store["chats"][chat_id]["title"] = message[:32] + "..." if len(message) > 32 else message

    session_messages = store["chats"][chat_id]["messages"]
    session_messages[0] = {"role": "system", "content": get_system_prompt(store["memories"])}
    steps = web_run_agent(message, session_messages, store["memories"])
    
    store["chats"][chat_id]["messages"] = session_messages
    save_store(store)
    
    return jsonify({"steps": steps, "chat_id": chat_id, "title": store["chats"][chat_id]["title"]})

@app.route("/api/chat/stream", methods=["POST"])
def chat_stream():
    """SSE streaming endpoint - yields steps in real-time."""
    data = request.get_json(force=True) or {}
    chat_id = data.get("chat_id")
    message = (data.get("message") or "").strip()

    if not message:
        def empty_gen():
            yield json.dumps({"type": "error", "text": "Empty command syntax query submission."}) + "\n"
        return Response(stream_with_context(empty_gen()), content_type="text/event-stream")

    store = load_store()

    if not chat_id or chat_id not in store["chats"]:
        chat_id = uuid.uuid4().hex
        store["chats"][chat_id] = {
            "title": message[:32] + "..." if len(message) > 32 else message,
            "messages": [{"role": "system", "content": get_system_prompt(store["memories"])}]
        }
    elif store["chats"][chat_id]["title"] == "Untitled Operations Thread":
        store["chats"][chat_id]["title"] = message[:32] + "..." if len(message) > 32 else message

    session_messages = store["chats"][chat_id]["messages"]
    session_messages[0] = {"role": "system", "content": get_system_prompt(store["memories"])}

    def generate():
        try:
            for step in web_run_agent_stream(message, session_messages, store["memories"]):
                yield json.dumps(step) + "\n"
            # Send done signal with metadata
            yield json.dumps({"type": "done", "chat_id": chat_id, "title": store["chats"][chat_id]["title"]}) + "\n"
            # Persist after streaming completes
            store["chats"][chat_id]["messages"] = session_messages
            save_store(store)
        except Exception as e:
            yield json.dumps({"type": "error", "text": str(e)}) + "\n"

    return Response(stream_with_context(generate()), content_type="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

# --- Document Editor AI Route ---
@app.route("/api/chat/doc", methods=["POST"])
def chat_doc():
    data = request.get_json(force=True) or {}
    instruction = (data.get("instruction") or "").strip()
    doc_id = (data.get("doc_id") or "").strip()

    if not instruction:
        return jsonify({"error": "No instruction provided."}), 400

    store = load_store()
    doc = store.get("documents", {}).get(doc_id)
    if not doc:
        return jsonify({"error": "Document not found."}), 404

    system_prompt = "You are an AI document editor. Edit the document based on the user's request.\n\n"
    system_prompt += "CURRENT DOCUMENT:\n---BEGIN DOCUMENT---\n"
    system_prompt += doc["content"]
    system_prompt += "\n---END DOCUMENT---\n\n"
    system_prompt += "INSTRUCTIONS:\n"
    system_prompt += "- Fix grammar, rewrite, add content, or generate text as requested.\n"
    system_prompt += "- Return the COMPLETE edited document with every single line.\n"
    system_prompt += "- NEVER return only the changes or a summary of the content.\n"
    system_prompt += "- Keep the original text exactly as-is unless the user asks to change it.\n\n"
    system_prompt += "OUTPUT FORMAT (follow exactly):\n"
    system_prompt += "SUMMARY: <one line describing what you did>\n"
    system_prompt += "Then immediately output the raw document text with NO blank line.\n\n"
    system_prompt += "ABSOLUTELY FORBIDDEN:\n"
    system_prompt += "- Never start extra sentences like 'Here's a draft', 'Here is', 'I have', 'I've', 'The following', 'Below is', 'This is'.\n"
    system_prompt += "- Never add any words addressed to the user. No 'you', 'your', 'I hope', 'let me know'.\n"
    system_prompt += "- Never wrap the content in code blocks, quotes, or markdown.\n"
    system_prompt += "- Output ONLY the SUMMARY line and then the raw text. Nothing else.\n\n"
    system_prompt += "Example CORRECT output:\n"
    system_prompt += "SUMMARY: Fixed grammar and spelling\n"
    system_prompt += "Hello World! This is my document.\n\n"
    system_prompt += "Example WRONG output:\n"
    system_prompt += "SUMMARY: Fixed grammar\n"
    system_prompt += "Here's the corrected version:\n"
    system_prompt += "Hello World! This is my document."

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": instruction}
    ]

    response = call_ollama(messages)
    if response is None:
        return jsonify({"error": "AI failed to respond. Check that Ollama is running."}), 500
    if not response.strip():
        response = "SUMMARY: No changes made."

    summary = ""
    body_lines = []
    lines = response.split('\n')
    found_summary = False
    for line in lines:
        if line.startswith("SUMMARY:"):
            summary = line.replace("SUMMARY:", "").strip()
            found_summary = True
        elif found_summary:
            body_lines.append(line)
        elif not found_summary and not line.strip():
            continue

    # If no SUMMARY: prefix found, the whole response is potential content
    if not found_summary:
        body_lines = lines[:]

    # Aggressively strip preamble lines before the actual document content
    while body_lines:
        line = body_lines[0].strip()
        if not line:
            body_lines.pop(0)
            continue
        lower = line.lower()
        # Check for known preamble patterns (AI talking to user instead of outputting content)
        starts_with_trigger = any(lower.startswith(t) for t in ("here's", "here is", "here are", "i have", "i've", "i am", "the following", "below is", "below are", "this is", "that is", "the edited", "the corrected", "the updated", "the revised", "edited document", "corrected document", "updated document", "please find", "attached is", "here you", "i hope", "let me"))
        contains_you = "you" in lower or "your" in lower
        is_short = len(line) < 100
        ends_with_colon = line.rstrip().endswith(":")
        # Strip if it's addressed to the user or starts with preamble triggers
        if starts_with_trigger or (is_short and ends_with_colon and contains_you):
            body_lines.pop(0)
        else:
            break
    # Remove leading blank lines and ``` markers
    body_lines = [l for l in body_lines if not l.strip().startswith("```")]
    while body_lines and not body_lines[0].strip():
        body_lines.pop(0)
    while body_lines and not body_lines[-1].strip():
        body_lines.pop()

    new_content = "\n".join(body_lines).strip()

    # Fallback: try extracting content between ``` blocks if raw body is empty
    if not new_content or len(new_content) < 5:
        in_block = False
        block_lines = []
        for line in lines:
            if line.strip().startswith("```"):
                if in_block:
                    new_content = "\n".join(block_lines).strip()
                    break
                in_block = True
                continue
            if in_block:
                block_lines.append(line)

    # Fallback: try content after CONTENT: line
    if not new_content or len(new_content) < 5:
        in_content = False
        content_lines = []
        for line in lines:
            if line.strip().startswith("CONTENT:"):
                in_content = True
                rest = line.replace("CONTENT:", "", 1).strip()
                if rest:
                    content_lines.append(rest)
                continue
            if in_content:
                content_lines.append(line)
        if content_lines:
            candidate = "\n".join(content_lines).strip()
            if len(candidate) > 5:
                new_content = candidate

    if not new_content:
        new_content = doc["content"]

    doc_write(doc_id, new_content)

    return jsonify({
        "summary": summary or "Document updated.",
        "new_content": new_content
    })

# --- Memory Routes ---
@app.route("/api/memories/add", methods=["POST"])
def add_memory():
    data = request.get_json(force=True) or {}
    text = (data.get("text") or "").strip()
    if not text:
        return jsonify({"error": "Empty string context"}), 400
    store = load_store()
    if text not in store["memories"]:
        store["memories"].append(text)
        save_store(store)
    return jsonify({"ok": True, "memories": store["memories"]})

@app.route("/api/memories/delete", methods=["POST"])
def delete_memory():
    data = request.get_json(force=True) or {}
    index = data.get("index")
    store = load_store()
    if index is not None and 0 <= index < len(store["memories"]):
        store["memories"].pop(index)
        save_store(store)
    return jsonify({"ok": True, "memories": store["memories"]})

# --- Document Workspaces Routes ---
@app.route("/api/documents/get", methods=["GET"])
def get_document():
    doc_id = request.args.get("doc_id")
    store = load_store()
    if doc_id not in store["documents"]:
        return jsonify({"error": "File object not found"}), 404
    return jsonify(store["documents"][doc_id])

@app.route("/api/documents/save", methods=["POST"])
def save_document():
    data = request.get_json(force=True) or {}
    doc_id = data.get("doc_id")
    title = data.get("title", "Untitled.txt").strip() or "Untitled.txt"
    content = data.get("content", "")
    
    store = load_store()
    if not doc_id:
        doc_id = uuid.uuid4().hex
        
    store["documents"][doc_id] = {"title": title, "content": content}
    save_store(store)
    return jsonify({"ok": True, "doc_id": doc_id, "title": title})

@app.route("/api/documents/new", methods=["POST"])
def new_document():
    store = load_store()
    doc_id = uuid.uuid4().hex
    store["documents"][doc_id] = {"title": "New_Buffer.txt", "content": ""}
    save_store(store)
    return jsonify({"doc_id": doc_id, "title": "New_Buffer.txt"})

@app.route("/api/documents/export", methods=["POST"])
def export_document():
    data = request.get_json(force=True) or {}
    doc_id = data.get("doc_id")
    filepath = data.get("filepath", "").strip()
    if not doc_id or not filepath:
        return jsonify({"error": "doc_id and filepath required"}), 400
    store = load_store()
    docs = store.get("documents", {})
    doc = docs.get(doc_id)
    if not doc:
        for did, d in docs.items():
            if d.get("title", "").lower() == doc_id.lower():
                doc = d
                break
    if not doc:
        return jsonify({"error": "Document not found"}), 404
    try:
        with open(filepath, "w") as f:
            f.write(doc["content"])
        return jsonify({"ok": True, "filepath": filepath})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/documents/import", methods=["POST"])
def import_document():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "Empty filename"}), 400
    try:
        content = file.read().decode("utf-8")
    except Exception:
        return jsonify({"error": "File must be text-based"}), 400
    store = load_store()
    doc_id = uuid.uuid4().hex
    store["documents"][doc_id] = {"title": file.filename, "content": content}
    save_store(store)
    return jsonify({"ok": True, "doc_id": doc_id, "title": file.filename})

# --- Calendar Reminder Routes ---
@app.route("/api/reminders/get", methods=["GET"])
def get_reminders():
    store = load_store()
    return jsonify({"reminders": store["reminders"]})

@app.route("/api/reminders/add", methods=["POST"])
def add_reminder():
    data = request.get_json(force=True) or {}
    date_key = (data.get("date") or "").strip()
    text = (data.get("text") or "").strip()
    if not date_key or not text:
        return jsonify({"error": "Both a date and reminder text are required"}), 400

    store = load_store()
    store["reminders"].setdefault(date_key, [])
    store["reminders"][date_key].append(text)
    save_store(store)
    return jsonify({"ok": True, "reminders": store["reminders"]})

@app.route("/api/reminders/delete", methods=["POST"])
def delete_reminder():
    data = request.get_json(force=True) or {}
    date_key = (data.get("date") or "").strip()
    index = data.get("index")

    store = load_store()
    entries = store["reminders"].get(date_key, [])
    if index is not None and 0 <= index < len(entries):
        entries.pop(index)
        if not entries:
            store["reminders"].pop(date_key, None)
        else:
            store["reminders"][date_key] = entries
        save_store(store)
    return jsonify({"ok": True, "reminders": store["reminders"]})

# --- Direct File Conversion Backend Engine ---
@app.route("/api/convert", methods=["POST"])
def convert_file():
    if "file" not in request.files or "target_format" not in request.form:
        return jsonify({"error": "Missing uploaded file chunk or programmatic format specification"}), 400

    file = request.files["file"]
    target_format = request.form["target_format"].strip().lower()
    filename = file.filename
    base_name, ext = os.path.splitext(filename)
    ext = ext.lower().replace(".", "")

    if not filename:
        return jsonify({"error": "No processing file designated"}), 400

    file_bytes = file.read()
    output_buffer = io.BytesIO()
    out_mimetype = "application/octet-stream"
    out_filename = f"{base_name}.{target_format}"

    try:
        # PDF -> TXT
        if ext == "pdf" and target_format == "txt":
            reader = PdfReader(io.BytesIO(file_bytes))
            text_content = []
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text_content.append(extracted)
            output_buffer.write("\n".join(text_content).encode("utf-8"))
            out_mimetype = "text/plain"

        # DOC/DOCX -> TXT
        elif ext in ["doc", "docx"] and target_format == "txt":
            doc = Document(io.BytesIO(file_bytes))
            text_content = [para.text for para in doc.paragraphs]
            output_buffer.write("\n".join(text_content).encode("utf-8"))
            out_mimetype = "text/plain"

        # DOCX -> PDF
        elif ext == "docx" and target_format == "pdf":
            return jsonify({"error": "DOCX to PDF conversion requires the 'fpdf2' library. Install it with: pip install fpdf2"}), 400

        # PDF -> DOCX
        elif ext == "pdf" and target_format == "docx":
            reader = PdfReader(io.BytesIO(file_bytes))
            doc = Document()
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
            doc.save(output_buffer)
            out_mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # IMAGE transformations (PNG -> JPEG, JPEG -> PNG)
        elif ext in ["png", "jpg", "jpeg"] and target_format in ["png", "jpeg", "jpg"]:
            img = Image.open(io.BytesIO(file_bytes))
            if target_format in ["jpeg", "jpg"] and img.mode in ("RGBA", "LA", "P"):
                if img.mode != "RGBA":
                    img = img.convert("RGBA")
                background = Image.new("RGB", img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[3])
                img = background
            
            fmt_key = "JPEG" if target_format in ["jpeg", "jpg"] else "PNG"
            img.save(output_buffer, format=fmt_key)
            normalized_format = "jpeg" if target_format in ["jpeg", "jpg"] else target_format
            out_mimetype = f"image/{normalized_format}"

        # IMAGE -> PDF (PNG/JPEG -> PDF)
        elif ext in ["png", "jpg", "jpeg"] and target_format == "pdf":
            img = Image.open(io.BytesIO(file_bytes))
            if img.mode in ("RGBA", "LA", "P"):
                if img.mode != "RGBA":
                    img = img.convert("RGBA")
                background = Image.new("RGB", img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[3])
                img = background
            img.save(output_buffer, format="PDF")
            out_mimetype = "application/pdf"

        # PDF -> IMAGE (Extracts first page render layer/fallback)
        elif ext == "pdf" and target_format in ["png", "jpeg", "jpg"]:
            # Fallback wrapper for raw bytes generation context alignment
            return jsonify({"error": "PDF to Image conversion requires host level poppler binaries. Direct stream rejected."}), 400

        else:
            return jsonify({"error": f"Unsupported dynamic execution pipeline matching: .{ext} to .{target_format}"}), 400

        output_buffer.seek(0)
        return send_file(
            output_buffer,
            mimetype=out_mimetype,
            as_attachment=True,
            download_name=out_filename
        )

    except Exception as e:
        return jsonify({"error": f"Conversion matrix failed runtime evaluation: {str(e)}"}), 500


PAGE_HTML = r"""
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>⬡ litelaw AI Workspace</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;500;600;700;800&display=swap" rel="stylesheet">
<style>
@font-face {
  font-family: 'JetBrainsMono Nerd Font';
  src: url('https://cdn.jsdelivr.net/gh/ryanoasis/nerd-fonts@v3.2.1/patched-fonts/JetBrainsMono/Ligatures/Regular/JetBrainsMonoNerdFontMono-Regular.ttf') format('truetype');
  font-weight: 400; font-style: normal; font-display: swap;
}
:root{
  --bg-0:#07040f;
  --bg-1:#0d0620;
  --bg-2:#160a33;
  --panel:#120a26cc;
  --panel-solid:#140b28;
  --sidebar-bg:#090514;
  --border:#3a2470;
  --border-soft:#2a1a52;
  --violet-1:#a78bfa;
  --violet-2:#8b5cf6;
  --violet-3:#7c3aed;
  --magenta:#e879f9;
  --cyan:#67e8f9;
  --text-0:#ede9fe;
  --text-1:#c4b5fd;
  --text-dim:#8a7cad;
  --green:#4ade80;
  --red:#f87171;
  --yellow:#facc15;
  --mono: 'JetBrainsMono Nerd Font', 'JetBrains Mono', ui-monospace, monospace;
  --ease-smooth: cubic-bezier(0.4, 0, 0.2, 1);
  --ease-spring: cubic-bezier(0.34, 1.56, 0.64, 1);
  --ease-bouncy: cubic-bezier(0.68, -0.55, 0.265, 1.55);
  --ease-out: cubic-bezier(0, 0, 0.2, 1);
  --duration-fast: 150ms;
  --duration-normal: 300ms;
  --duration-slow: 450ms;
}

/* KEYFRAMES FOR ANIMATIONS */
@keyframes popInBouncy {
  0% { opacity: 0; transform: scale(0.9) translateY(10px); }
  60% { opacity: 1; transform: scale(1.03) translateY(-2px); }
  100% { opacity: 1; transform: scale(1) translateY(0); }
}
@keyframes slideUpSmooth {
  0% { opacity: 0; transform: translateY(20px); }
  100% { opacity: 1; transform: translateY(0); }
}
@keyframes slideInRightSmooth {
  0% { opacity: 0; transform: translateX(-15px); }
  100% { opacity: 1; transform: translateX(0); }
}
@keyframes fadeScaleIn {
  0% { opacity: 0; transform: scale(0.97); }
  100% { opacity: 1; transform: scale(1); }
}
@keyframes pulseBouncy {
   0% { transform: scale(1); }
   50% { transform: scale(1.05); }
   100% { transform: scale(1); }
}

@Keyframes smoothSpin {
   0% { transform: rotate(0deg) scale(1); }
   50% { transform: rotate(180deg) scale(1.05); }
   100% { transform: rotate(360deg) scale(1); }
}

@Keyframes gentlePulse {
   0%, 100% { opacity: 1; transform: scale(1); }
   50% { opacity: 0.85; transform: scale(1.02); }
}

{ box-sizing:border-box; }
html,body{
  height:100%; margin:0; font-family: var(--mono);
  background: var(--bg-0); color: var(--text-0); overflow:hidden;
}
.nebula{
  position:fixed; inset:0; z-index:0; pointer-events:none;
  background:
    radial-gradient(ellipse 900px 600px at 15% 15%, rgba(139,92,246,0.12), transparent 60%),
    radial-gradient(ellipse 800px 700px at 85% 25%, rgba(232,121,249,0.08), transparent 60%),
    linear-gradient(180deg, var(--bg-1) 0%, var(--bg-0) 100%);
}
#stars{ position:fixed; inset:0; z-index:0; pointer-events:none; }

.app-frame{
  position:relative; z-index:1; height:100vh;
  display:flex; flex-direction:column;
}
header{
  display:flex; align-items:center; justify-content:space-between;
  padding:12px 20px; border-bottom:1px solid var(--border-soft);
  background: rgba(20,11,40,0.85); backdrop-filter: blur(12px);
  animation: fadeScaleIn 0.8s var(--ease-smooth);
}
.brand{ display:flex; align-items:center; gap:12px; }
.brand-mark{
  width:32px; height:32px; border-radius:8px;
  display:flex; align-items:center; justify-content:center;
  background: radial-gradient(circle at 30% 30%, var(--violet-1), var(--violet-3) 70%);
  box-shadow: 0 0 14px rgba(139,92,246,0.4); font-size: 18px;;
  transition: transform 0.4s var(--ease-spring);
}
.brand-mark:hover {
  transform: rotate(15deg) scale(1.08);
}
.brand-text h1{
  margin:0; font-size: 20px;; font-weight:800; letter-spacing:0.5px;
  background: linear-gradient(90deg, var(--violet-1), var(--magenta));
  -webkit-background-clip:text; background-clip:text; color:transparent;
}
.brand-text span{ display:block; font-size: 14px;; color:var(--text-dim); }

.workspace-layout{ display:flex; flex:1; overflow:hidden; }

/* --- SIDEBAR PANEL --- */
.sidebar{
  width:280px; background: var(--sidebar-bg);
  border-right:1px solid var(--border-soft);
  display:flex; flex-direction:column; gap:20px; padding:16px; overflow-y:auto;
  animation: slideInRightSmooth 0.5s var(--ease-smooth);
}
.sidebar-section-title{
  font-size: 10.5px; text-transform:uppercase; color:var(--text-dim);
  letter-spacing:1px; margin-bottom:8px; font-weight:700;
  display:flex; align-items:center; justify-content:space-between;
}
.sidebar-btn{
  width:100%; font-family:var(--mono); text-align:left;
  background:rgba(139,92,246,0.06); border:1px solid var(--border-soft);
  color:var(--text-1); padding:9px 12px; border-radius:8px;
  font-size: 12.5px; cursor:pointer; transition: all var(--duration-fast) var(--ease-spring);
  display:flex; align-items:center; gap:8px; margin-bottom:6px;
  position:relative; overflow:hidden;
  transform-origin: center;
}
.sidebar-btn::before {
  content: '';
  position: absolute;
  top: 0;
  left: -100%;
  width: 100%;
  height: 100%;
  background: linear-gradient(90deg, transparent, rgba(139,92,246,0.1), transparent);
  transition: left var(--duration-slow) var(--ease-smooth);
}
.sidebar-btn:hover::before, .sidebar-btn.active::before {
  left: 100%;
}
.sidebar-btn:hover, .sidebar-btn.active{
  background:rgba(139,92,246,0.16); 
  border-color: var(--violet-1); 
  color:#fff;
  transform: translateY(-2px);
  box-shadow: 0 4px 12px rgba(139,92,246,0.3);
}
.sidebar-btn:active {
  transform: scale(0.96) translateY(0);
  box-shadow: 0 1px 4px rgba(139,92,246,0.2);
}

.history-list, .doc-list { 
  display: flex; 
  flex-direction: column; 
  gap: 4px; 
  max-height: 160px;
  width: 100%; 
  overflow-y: auto;
  overflow-x: hidden; 
  -ms-overflow-style: auto;
  scrollbar-width: thin;
  scrollbar-color: rgba(167,139,250,0.35) transparent;
}

.history-list::-webkit-scrollbar,
.doc-list::-webkit-scrollbar {
  width: 6px;
}
.history-list::-webkit-scrollbar-track,
.doc-list::-webkit-scrollbar-track {
  background: transparent;
}
.history-list::-webkit-scrollbar-thumb,
.doc-list::-webkit-scrollbar-thumb {
  background: rgba(167,139,250,0.35);
  border-radius: 8px;
}
.history-list::-webkit-scrollbar-thumb:hover,
.doc-list::-webkit-scrollbar-thumb:hover {
  background: rgba(167,139,250,0.6);
}
.list-item-link {
  display: block;
  width: 100%;
  flex-shrink: 0;
  padding: 6px 10px; 
    font-size: 12px; 
  border-radius: 6px; 
  cursor: pointer;
  white-space: nowrap; 
  overflow: hidden; 
  text-overflow: ellipsis;
  color: var(--text-dim); 
  transition: all var(--duration-fast) var(--ease-spring);
  animation: slideInRightSmooth 0.4s var(--ease-smooth) backwards;
}

.list-item-link:nth-child(1) { animation-delay: 0.05s; }
.list-item-link:nth-child(2) { animation-delay: 0.1s; }
.list-item-link:nth-child(3) { animation-delay: 0.15s; }
.list-item-link:nth-child(4) { animation-delay: 0.2s; }
.list-item-link:nth-child(5) { animation-delay: 0.25s; }

.list-item-link:hover, .list-item-link.active {
  background: rgba(167,139,250,0.08); 
  color: var(--text-0);
  transform: translateX(3px);
}
.list-item-link:active {
  transform: scale(0.98);
}
/* --- MAIN INTERACTIVE VIEW AREA --- */
.main-stage{ flex:1; display:flex; flex-direction:column; background:rgba(7,4,15,0.4); overflow:hidden; }
.stage-panel{ 
    display:none; 
    flex:1; 
    flex-direction:column; 
    overflow:hidden;
    opacity:0;
    transform: translateY(15px);
    transition: opacity var(--duration-normal) var(--ease-out), transform var(--duration-normal) var(--ease-out);
  }
  .stage-panel.active{ 
    display:flex; 
    min-height:0;
    opacity:1;
    transform: translateY(0);
    animation: slideUpSmooth 0.4s var(--ease-spring);
  }

/* CHAT PANEL SCROLLER */
.chat-scroller { 
  flex: 1; 
  min-height: 0;
  width: 100%;
  overflow-y: auto;
  overflow-x: hidden;
  padding: 24px 0; 
  display: flex;
  flex-direction: column;
  position: relative;
  -ms-overflow-style: auto;
  scrollbar-width: thin;
  scrollbar-color: rgba(167,139,250,0.35) transparent;
  scroll-behavior: smooth;
}
.chat-scroller::-webkit-scrollbar {
  width: 8px;
  background: transparent;
  transition: background var(--duration-fast) var(--ease-smooth);
}
.chat-scroller::-webkit-scrollbar-thumb {
  background: rgba(167,139,250,0.35);
  border-radius: 8px;
  transition: background var(--duration-fast) var(--ease-smooth), transform var(--duration-fast) var(--ease-spring);
}
.chat-scroller::-webkit-scrollbar-thumb:hover {
  background: rgba(167,139,250,0.6);
  transform: scale(1.1);
}

.scroll-to-bottom-btn {
  position: absolute;
  bottom: 16px;
  left: 50%;
  transform: translateX(-50%) translateY(20px);
  display: none;
  align-items: center;
  gap: 6px;
  padding: 7px 14px;
  border-radius: 999px;
  border: 1px solid var(--border-soft);
  background: rgba(20,11,40,0.92);
  backdrop-filter: blur(8px);
  color: var(--text-1);
  font-family: var(--mono);
  font-size: 15px;;
  cursor: pointer;
  box-shadow: 0 4px 14px rgba(0,0,0,0.35);
  z-index: 5;
  transition: all 0.3s var(--ease-spring);
  opacity: 0;
}
.scroll-to-bottom-btn.visible { 
  display: flex; 
  transform: translateX(-50%) translateY(0);
  opacity: 1;
}
.scroll-to-bottom-btn:hover { background: rgba(139,92,246,0.18); transform: translateX(-50%) translateY(-2px); }
.scroll-to-bottom-btn:active { transform: translateX(-50%) translateY(1px) scale(0.95); }

/* BUBBLES CONTAINER */
.chat-container { 
  width: 100%;
  max-width: 800px; 
  margin: 0 auto; 
  padding: 0 20px; 
  display: flex; 
  flex-direction: column; 
  gap: 16px;
  flex-grow: 1;
}
.row{ display:flex; width:100%; }
.row.user{ justify-content:flex-end; }
.row.assistant{ justify-content:flex-start; }
.bubble{ 
  max-width:82%; padding:11px 15px; border-radius:12px; font-size: 15px;; line-height:1.5; white-space:pre-wrap; 
  animation: popInBouncy 0.5s var(--ease-bouncy) backwards;
  transform-origin: left bottom;
}
.row.user .bubble {
  transform-origin: right bottom;
}
.bubble.user{
  background: linear-gradient(135deg, var(--violet-3), #5b21b6); color:#fff;
  border: 1px solid rgba(167,139,250,0.3); border-bottom-right-radius:2px;
}
.bubble.final{ background: var(--panel-solid); border:1px solid var(--border-soft); border-bottom-left-radius:2px; }
.bubble.final::before{ content:"⬡ litelaw response"; display:block; font-size: 12px;; color:var(--violet-1); margin-bottom:4px; font-weight:700; }
.thought{ max-width:82%; font-size: 13.5px;; color:var(--text-dim); font-style:italic; padding:6px 10px; min-height:24px; animation: slideUpSmooth 0.4s var(--ease-smooth) backwards; }
.think-dots{ display:inline-flex; gap:4px; padding:4px 0; }
.think-dots span{ width:8px; height:8px; border-radius:50%; background:var(--violet-1); animation: dotPulse 1.4s infinite; opacity:0.3; font-style:normal; font-size:0; }
.think-dots span:nth-child(2){ animation-delay:0.2s; }
.think-dots span:nth-child(3){ animation-delay:0.4s; }
@keyframes dotPulse{ 0%,80%,100%{ opacity:0.2; transform:scale(0.8); } 40%{ opacity:1; transform:scale(1.2); } }
.msg-time{ font-size:10px; color:var(--text-dim); padding:0 4px; opacity:0.7; }
.msg-time-inline{ display:block; font-size:10px; color:rgba(255,255,255,0.4); margin-top:6px; text-align:right; }

.term-block{ width:82%; border-radius:8px; overflow:hidden; border:1px solid var(--border-soft); background:#05020a; animation: slideInRightSmooth 0.4s var(--ease-smooth) backwards; }
.term-head{ display:flex; align-items:center; gap:6px; padding:6px 12px; background:rgba(139,92,246,0.06); font-size: 12px;; color:var(--text-dim); }
.term-dots span{ width:6px; height:6px; border-radius:50%; display:inline-block; background:#3a2470; transition: background 0.3s; }
.term-block:hover .term-dots span:nth-child(1) { background: var(--red); }
.term-block:hover .term-dots span:nth-child(2) { background: var(--yellow); }
.term-block:hover .term-dots span:nth-child(3) { background: var(--green); }

.term-body{ padding:10px 12px; font-size: 14px;; white-space:pre-wrap; }
.term-body.command{ color:var(--cyan); }
.term-body.command::before{ content:"$ "; color:var(--violet-1); }
.term-body.output{ color:var(--text-1); }

/* --- MULTI-AGENT: TASK QUEUE / AGENT BADGES / REVIEW --- */
.task-queue-block{ width:82%; border-radius:8px; overflow:hidden; border:1px solid var(--border-soft); background:#05020a; animation: slideUpSmooth 0.4s var(--ease-smooth) backwards; }
.task-queue-head{
  display:flex; align-items:center; gap:6px; padding:6px 12px; background:rgba(139,92,246,0.08);
  font-size: 12px;; color:var(--violet-1); text-transform:uppercase; letter-spacing:0.8px; font-weight:700;
}
.task-queue-body{ padding:10px 12px; display:flex; flex-direction:column; gap:7px; }
.queue-item{ display:flex; align-items:center; gap:8px; font-size: 14px;; color:var(--text-dim); animation: slideInRightSmooth 0.3s var(--ease-smooth) backwards; }
.queue-item:nth-child(1) { animation-delay: 0.1s; }
.queue-item:nth-child(2) { animation-delay: 0.2s; }
.queue-item:nth-child(3) { animation-delay: 0.3s; }
.queue-item:nth-child(4) { animation-delay: 0.4s; }
.queue-item .q-icon{ width:14px; text-align:center; flex-shrink:0; }
.queue-item .q-agent{
  font-size: 11px;; text-transform:uppercase; padding:1px 6px; border-radius:4px; flex-shrink:0;
  background:rgba(139,92,246,0.15); color:var(--violet-1); letter-spacing:0.4px;
}
.queue-item .q-text{ flex:1; }
.queue-item.active{ color:var(--cyan); }
.queue-item.active .q-icon{ animation: queuePulse 1s infinite; color:var(--cyan); }
.queue-item.done{ color:var(--green); transition: all 0.3s ease; }
.queue-item.done .q-icon{ color:var(--green); }
.queue-item.done .q-text{ text-decoration:line-through; opacity:0.7; }
.queue-item.failed{ color:var(--red); }
.queue-item.failed .q-icon{ color:var(--red); }
@keyframes queuePulse{ 0%,100%{ opacity:1; transform: scale(1); } 50%{ opacity:0.25; transform: scale(0.9); } }

.agent-badge{
  display:inline-block; font-size: 11px;; text-transform:uppercase; letter-spacing:0.5px;
  padding:1.5px 6px; border-radius:4px; margin-right:6px; vertical-align:middle;
  background:rgba(139,92,246,0.15); color:var(--violet-1); font-weight:700;
}
.agent-badge.agent-researcher{ background:rgba(103,232,249,0.15); color:var(--cyan); }
.agent-badge.agent-coder{ background:rgba(250,204,21,0.15); color:var(--yellow); }
.agent-badge.agent-executor{ background:rgba(232,121,249,0.15); color:var(--magenta); }
.agent-badge.agent-reviewer{ background:rgba(74,222,128,0.15); color:var(--green); }
.agent-badge.agent-planner{ background:rgba(167,139,250,0.15); color:var(--violet-1); }

.review-block{ width:82%; border-radius:10px; padding:11px 15px; font-size: 15px;; line-height:1.5; border:1px solid var(--border-soft); background:var(--panel-solid); white-space:pre-wrap; animation: popInBouncy 0.4s var(--ease-spring) backwards;}
.review-block.pass{ border-color: rgba(74,222,128,0.45); }
.review-block.pass::before{ content:"✅ Reviewer verdict: PASS"; display:block; font-size: 12px;; color:var(--green); margin-bottom:4px; font-weight:700; }
.review-block.fail{ border-color: rgba(248,113,113,0.45); }
.review-block.fail::before{ content:"⚠ Reviewer verdict: FAIL"; display:block; font-size: 12px;; color:var(--red); margin-bottom:4px; font-weight:700; }

/* MEMORY PANEL */
.memory-workspace, .editor-workspace, .converter-workspace{ max-width:800px; width:100%; margin:30px auto; padding:0 20px; display:flex; flex-direction:column; gap:20px; }
.memory-input-group{ display:flex; gap:10px; }
.input-field{
  flex:1; background:rgba(139,92,246,0.06); border:1px solid var(--border);
  border-radius:8px; padding:10px 14px; color:#fff; font-family:var(--mono); font-size: 15px;;
  transition: border-color 0.3s ease, box-shadow 0.3s ease;
}
.input-field:focus {
  border-color: var(--violet-1);
  box-shadow: 0 0 8px rgba(139,92,246,0.3);
  outline: none;
}
.memory-card-list{ display:flex; flex-direction:column; gap:10px; margin-top:10px; }
.memory-item{
  background:var(--panel-solid); border:1px solid var(--border-soft);
  padding:12px 16px; border-radius:8px; display:flex; align-items:center; justify-content:space-between; font-size: 14.5px;;
  animation: fadeScaleIn 0.3s var(--ease-smooth) backwards;
  transition: transform 0.2s var(--ease-spring), box-shadow 0.2s var(--ease-smooth);
}
.memory-item:hover {
  transform: translateY(-2px);
  box-shadow: 0 4px 12px rgba(0,0,0,0.3);
}
.delete-btn{ background:transparent; border:none; color:var(--red); cursor:pointer; font-family:var(--mono); transition: transform 0.2s; }
.delete-btn:hover { transform: scale(1.1); }
.delete-btn:active { transform: scale(0.9); }

/* EDITOR WORKSPACE */
.editor-workspace{
  max-width:1000px; width:100%; margin:20px auto; padding:0 20px;
  display:flex; flex-direction:column; gap:12px; flex:1; min-height:0;
}
.editor-meta{ display:flex; gap:12px; width:100%; flex-wrap:wrap; }
.editor-body{ display:flex; flex:1; min-height:0; gap:12px; }
.editor-textarea{
  flex:1; min-height:200px; background: #05020a; border:1px solid var(--border-soft);
  border-radius:8px; padding:16px; color:var(--text-0); font-family:var(--mono); font-size: 15px;; line-height:1.6; resize:none;
  transition: border-color 0.3s ease, box-shadow 0.3s ease;
}
.editor-textarea:focus {
  border-color: var(--violet-2);
  box-shadow: inset 0 0 8px rgba(139,92,246,0.1);
  outline: none;
}
.editor-ai-panel{
  width:340px; flex-shrink:0; display:flex; flex-direction:column;
  background:var(--panel-solid); border:1px solid var(--border-soft);
  border-radius:10px; padding:12px; min-height:0;
  animation: slideInRightSmooth 0.5s var(--ease-spring);
}
.editor-ai-header{ font-size: 13px;; font-weight:700; color:var(--violet-1); margin-bottom:8px; flex-shrink:0; text-transform:uppercase; letter-spacing:0.5px; }
.editor-ai-messages{
  flex:1; overflow-y:auto; display:flex; flex-direction:column; gap:6px;
  padding-right:4px; margin-bottom:8px; min-height:0;
  -ms-overflow-style:auto; scrollbar-width:thin;
  scrollbar-color:rgba(167,139,250,0.35) transparent;
}
.editor-ai-messages::-webkit-scrollbar{ width:6px; }
.editor-ai-messages::-webkit-scrollbar-track{ background:transparent; }
.editor-ai-messages::-webkit-scrollbar-thumb{ background:rgba(167,139,250,0.35); border-radius:8px; }
.editor-ai-messages::-webkit-scrollbar-thumb:hover{ background:rgba(167,139,250,0.6); }
.editor-ai-welcome{ font-size: 13px;; color:var(--text-dim); font-style:italic; padding:4px 0; }
.editor-ai-msg{ font-size: 14px;; line-height:1.5; padding:6px 10px; border-radius:6px; white-space:pre-wrap; animation: popInBouncy 0.3s var(--ease-spring) backwards; transform-origin: left bottom;}
.editor-ai-msg.user{ background:rgba(139,92,246,0.12); color:var(--text-0); align-self:flex-end; transform-origin: right bottom;}
.editor-ai-msg.assistant{ background:rgba(7,4,15,0.5); border:1px solid var(--border-soft); color:var(--text-1); }
.editor-ai-msg.assistant.thought{ font-style:italic; color:var(--text-dim); font-size: 13px;; background:transparent; border:none; }
.editor-ai-msg.assistant.command{ color:var(--cyan); font-family:var(--mono); font-size: 13px;; }
.editor-ai-msg.assistant.output{ color:var(--text-1); font-family:var(--mono); font-size: 13px;; }
.editor-ai-msg.assistant.error{ color:var(--red); border-color:rgba(248,113,113,0.3); }
.editor-ai-msg.assistant.final{ color:var(--text-0); border-color:var(--violet-1); }
.editor-ai-input-row{
  display:flex; align-items:center; gap:8px; flex-shrink:0;
  background:rgba(139,92,246,0.04); border:1px solid var(--border);
  border-radius:8px; padding:6px 10px;
  transition: border-color 0.3s ease;
}
.editor-ai-input-row:focus-within { border-color: var(--violet-1); }
.editor-ai-input-row textarea{
  flex:1; resize:none; background:transparent; border:none; outline:none;
  color:#fff; font-family:var(--mono); font-size: 14px;;
}
.editor-ai-input-row button{
  width:28px; height:28px; border-radius:6px; border:none; cursor:pointer;
  background:linear-gradient(135deg, var(--violet-1), var(--violet-3));
  color:#fff; display:flex; align-items:center; justify-content:center; font-size: 15px;; flex-shrink:0;
  transition: transform 0.2s var(--ease-bouncy), box-shadow 0.2s;
}
.editor-ai-input-row button:hover { transform: scale(1.1); box-shadow: 0 2px 8px rgba(139,92,246,0.4); }
.editor-ai-input-row button:active { transform: scale(0.9); }

/* FILE CONVERTER STYLES */
.dropzone {
  border: 2px dashed var(--border);
  background: rgba(139,92,246,0.03);
  border-radius: 12px;
  padding: 40px 20px;
  text-align: center;
  cursor: pointer;
  transition: all 0.3s var(--ease-spring);
  display: flex;
  flex-direction: column;
  align-items: center;
  gap: 12px;
}
.dropzone:hover {
  background: rgba(139,92,246,0.08);
  border-color: var(--violet-1);
}
.dropzone.dragover {
  border-color: var(--cyan);
  background: rgba(103,232,249,0.08);
  box-shadow: 0 0 15px rgba(103,232,249,0.2);
  transform: scale(1.02);
}
.dropzone-icon {
  font-size: 34px;;
  color: var(--violet-1);
  transition: transform 0.3s var(--ease-spring);
}
.dropzone:hover .dropzone-icon { transform: translateY(-5px) scale(1.1); }
.dropzone-text {
  font-size: 15px;;
  color: var(--text-1);
}
.dropzone-hint {
  font-size: 15px;;
  color: var(--text-dim);
}
.conversion-matrix-box {
  background: var(--panel-solid);
  border: 1px solid var(--border-soft);
  border-radius: 10px;
  padding: 16px;
  display: flex;
  flex-direction: column;
  gap: 12px;
  animation: slideUpSmooth 0.4s var(--ease-smooth);
}
.select-wrapper {
  display: flex;
  align-items: center;
  gap: 12px;
  font-size: 15px;;
}
.custom-select {
  background: #05020a;
  border: 1px solid var(--border);
  color: var(--text-0);
  font-family: var(--mono);
  font-size: 15px;;
  padding: 8px 12px;
  border-radius: 6px;
  outline: none;
  cursor: pointer;
  transition: border-color 0.2s, box-shadow 0.2s;
}
.custom-select:focus {
  border-color: var(--violet-1);
  box-shadow: 0 0 8px rgba(139,92,246,0.2);
}
#selectedFileBanner {
  font-size: 14px;;
  color: var(--cyan);
  background: rgba(103,232,249,0.05);
  padding: 8px 12px;
  border-radius: 6px;
  border: 1px solid rgba(103,232,249,0.2);
  display: none;
  animation: fadeScaleIn 0.3s var(--ease-spring);
}

/* CHAT CONTROLS WRAPPER */
.input-wrap{ padding:14px 20px; background:rgba(7,4,15,0.4); }
.input-inner{ max-width:800px; margin:0 auto; display:flex; align-items:center; gap:10px; background:rgba(139,92,246,0.04); border:1px solid var(--border); border-radius:10px; padding:8px 12px; transition: border-color 0.3s, box-shadow 0.3s; }
.input-inner:focus-within { border-color: var(--violet-1); box-shadow: 0 0 10px rgba(139,92,246,0.15); }
textarea#msg{ flex:1; resize:none; background:transparent; border:none; outline:none; color:#fff; font-family:var(--mono); font-size: 15px;; }
.send-btn{ width:34px; height:34px; border-radius:8px; border:none; cursor:pointer; background:linear-gradient(135deg, var(--violet-1), var(--violet-3)); color:#fff; display:flex; align-items:center; justify-content:center; transition: transform 0.2s var(--ease-bouncy), box-shadow 0.2s; }
.send-btn:hover { transform: scale(1.1); box-shadow: 0 4px 12px rgba(139,92,246,0.4); }
.send-btn:active { transform: scale(0.9); }
.send-btn:disabled { filter: grayscale(1); opacity: 0.5; transform: none; cursor: not-allowed; }

.empty-state{ margin-top:10vh; text-align:center; color:var(--text-dim); animation: fadeScaleIn 0.8s var(--ease-smooth); }
.empty-state h2{ color:#fff; font-size: 18px;; margin-bottom:4px; }

/* CALCULATOR TOOL PANEL */
.tool-workspace{ max-width:420px; width:100%; margin:30px auto; padding:0 20px; display:flex; flex-direction:column; gap:16px; }
.calc-display{
  background:#05020a; border:1px solid var(--border-soft); border-radius:10px;
  padding:20px 16px; text-align:right; font-family:var(--mono);
  box-shadow: inset 0 2px 10px rgba(0,0,0,0.5);
}
.calc-display .calc-expr{ font-size: 14px;; color:var(--text-dim); min-height:16px; overflow-x:auto; white-space:nowrap; }
.calc-display .calc-value{ font-size: 32px;; color:#fff; font-weight:700; overflow-x:auto; white-space:nowrap; transition: font-size 0.1s ease; }
.calc-grid{ display:grid; grid-template-columns:repeat(4, 1fr); gap:10px; }
.calc-btn{
  font-family:var(--mono); font-size: 17px;; font-weight:600; cursor:pointer;
  padding:16px 0; border-radius:10px; border:1px solid var(--border-soft);
  background:rgba(139,92,246,0.06); color:var(--text-0); transition: all 0.15s var(--ease-bouncy);
}
.calc-btn:hover{ background:rgba(139,92,246,0.18); border-color:var(--violet-1); transform: translateY(-2px); box-shadow: 0 4px 10px rgba(139,92,246,0.2); }
.calc-btn:active{ transform: scale(0.92); box-shadow: none; background:rgba(139,92,246,0.3); }
.calc-btn.op{ color:var(--violet-1); background:rgba(139,92,246,0.1); }
.calc-btn.eq{ background:linear-gradient(135deg, var(--violet-1), var(--violet-3)); color:#fff; border:none; }
.calc-btn.eq:hover{ box-shadow: 0 4px 15px rgba(139,92,246,0.5); }
.calc-btn.clear{ color:var(--red); }
.calc-btn.wide{ grid-column: span 2; }

/* CALENDAR TOOL PANEL */
.calendar-workspace{ max-width:520px; width:100%; margin:30px auto; padding:0 20px; display:flex; flex-direction:column; gap:16px; }
.cal-header{ display:flex; align-items:center; justify-content:space-between; animation: slideUpSmooth 0.3s var(--ease-smooth); }
.cal-header h3{ margin:0; font-size: 17px;; color:#fff; transition: transform 0.3s var(--ease-smooth); }
.cal-nav-btn{
  width:30px; height:30px; border-radius:8px; border:1px solid var(--border-soft);
  background:rgba(139,92,246,0.06); color:var(--text-1); cursor:pointer; font-family:var(--mono);
  transition: all 0.2s var(--ease-bouncy);
}
.cal-nav-btn:hover{ background:rgba(139,92,246,0.18); border-color:var(--violet-1); transform: scale(1.1); }
.cal-nav-btn:active{ transform: scale(0.9); }
.cal-grid{ display:grid; grid-template-columns:repeat(7, 1fr); gap:6px; }
.cal-dow{ text-align:center; font-size: 12px;; color:var(--text-dim); text-transform:uppercase; letter-spacing:0.5px; padding-bottom:4px; }
.cal-day{
  position:relative; text-align:center; padding:10px 0; border-radius:8px; font-size: 14.5px;;
  border:1px solid transparent; color:var(--text-1); cursor:pointer;
  transition: all 0.2s var(--ease-spring);
  animation: fadeScaleIn 0.4s var(--ease-smooth) backwards;
}
.cal-day.empty{ visibility:hidden; cursor:default; }
.cal-day.today{ background:linear-gradient(135deg, var(--violet-1), var(--violet-3)); color:#fff; font-weight:700; box-shadow: 0 2px 8px rgba(139,92,246,0.4); }
.cal-day.selected{ border-color:var(--violet-1); background:rgba(139,92,246,0.14); transform: scale(1.05); }
.cal-day:not(.empty):not(.today):hover{ background:rgba(139,92,246,0.1); border-color:var(--border-soft); transform: translateY(-2px); }
.cal-day:not(.empty):active{ transform: scale(0.95); }
.cal-day.has-reminder::after{
  content:""; position:absolute; bottom:3px; left:50%; transform:translateX(-50%);
  width:4px; height:4px; border-radius:50%; background:var(--magenta);
}
.cal-day.today.has-reminder::after{ background:#fff; }
.cal-today-label{ font-size: 13.5px;; color:var(--text-dim); text-align:center; }

/* REMINDER SUB-PANEL */
.reminder-panel{
  background:var(--panel-solid); border:1px solid var(--border-soft); border-radius:10px;
  padding:14px 16px; display:flex; flex-direction:column; gap:10px;
  animation: slideUpSmooth 0.5s var(--ease-spring);
}
.reminder-panel-title{ font-size: 14.5px;; color:#fff; font-weight:700; transition: color 0.3s; }
.reminder-input-group{ display:flex; gap:8px; }
.reminder-input-group input{ flex:1; }
.reminder-item-list{ display:flex; flex-direction:column; gap:6px; }
.reminder-item{
  display:flex; align-items:center; justify-content:space-between; gap:8px;
  background:rgba(139,92,246,0.06); border:1px solid var(--border-soft); border-radius:6px;
  padding:8px 10px; font-size: 14px;; color:var(--text-1);
  animation: slideInRightSmooth 0.3s var(--ease-spring) backwards;
}
.reminder-item-empty{ font-size: 13.5px;; color:var(--text-dim); font-style:italic; }
.reminder-notify-banner{
  display:flex; align-items:flex-start; gap:8px; background:rgba(232,121,249,0.1);
  border:1px solid var(--magenta); border-radius:8px; padding:10px 12px; font-size: 14px;; color:var(--text-0);
  animation: popInBouncy 0.4s var(--ease-spring);
}
</style>
</head>
<body>
<div class="nebula"></div>
<canvas id="stars"></canvas>

<div class="app-frame">
  <header>
    <div class="brand">
      <div class="brand-mark">⬡</div>
      <div class="brand-text">
        <h1>litelaw</h1>
        <span>automation environment & control agent</span>
      </div>
    </div>
    <div style="display:flex; align-items:center; gap:10px; font-size: 13px;;">
      <span style="display:inline-block; width:6px; height:6px; border-radius:50%; background:var(--green); animation: pulseBouncy 2s infinite;"></span><span style="color:var(--green);">agent synchronized</span>
    </div>
  </header>

  <div class="workspace-layout">
    <div class="sidebar">
      <button class="sidebar-btn" id="newChatBtn">➔ + New Context Thread</button>
      
      <div>
        <div class="sidebar-section-title">Tools</div>
        <button class="sidebar-btn" id="memVaultTabLink">✦ Long-Term Memory</button>
        <button class="sidebar-btn" id="docEditorTabLink">📝 Document Editor</button>
        <button class="sidebar-btn" id="fileConverterTabLink">🔄 File Converter</button>
        <button class="sidebar-btn" id="calculatorTabLink">🧮 Calculator</button>
        <button class="sidebar-btn" id="calendarTabLink">📅 Calendar</button>
      </div>

      <div>
        <div class="sidebar-section-title">Execution History</div>
        <div class="history-list" id="historyContainer"></div>
      </div>

      <div>
        <div class="sidebar-section-title">Saved Buffers</div>
        <div class="doc-list" id="savedDocsContainer"></div>
      </div>
    </div>

    <div class="main-stage">
      
      <div class="stage-panel active" id="chatPanel">
        <div class="chat-scroller" id="chatScroller">
          <div class="chat-container" id="chatContainer">
            <div class="empty-state" id="emptyState">
              <h2>Autonomous Terminal Workspace</h2>
              <p style="font-size: 14px;;">Provide goals in standard plaintext. The engine compiles and acts locally.</p>
            </div>
          </div>
          <button class="scroll-to-bottom-btn" id="scrollToBottomBtn">↓ Jump to latest</button>
        </div>
        <div class="input-wrap">
          <div class="input-inner" style="position:relative;">
            <textarea id="msg" rows="1" placeholder="Instruct agent..." style="flex:1;"></textarea>
            <button class="send-btn" id="sendBtn">➤</button>
          </div>
        </div>
      </div>

      <div class="stage-panel" id="memoryPanel">
        <div class="memory-workspace">
          <h3>🧠 Core Long-Term Memory Guard</h3>
          <p style="font-size: 14px;; color:var(--text-dim); margin:0 0 10px 0;">Statements managed here persist across engine restarts and lock rules or preferences natively into the compilation logic loop.</p>
          <div class="memory-input-group">
            <input type="text" id="memoryInput" class="input-field" placeholder="Append context constraint rule...">
            <button class="sidebar-btn" id="addMemoryBtn" style="width:auto; margin:0;">Add Memory</button>
          </div>
          <div class="memory-card-list" id="memoryCardList"></div>
        </div>
      </div>

      <div class="stage-panel" id="editorPanel">
        <div class="editor-workspace">
          <h3>📝 Dynamic Document Buffer Workspace</h3>
          <div class="editor-meta">
            <input type="text" id="docTitle" class="input-field" placeholder="File name context (e.g. log.txt)">
            <button class="sidebar-btn" id="saveDocBtn" style="width:auto; margin:0;">💾 Commit Save File</button>
            <button class="sidebar-btn" id="newDocBtn" style="width:auto; margin:0; background:rgba(255,255,255,0.04)">+ New Doc</button>
            <button class="sidebar-btn" id="exportDocBtn" style="width:auto; margin:0;">⬇ Export</button>
            <button class="sidebar-btn" id="importDocBtn" style="width:auto; margin:0;">⬆ Import</button>
            <input type="file" id="importFileInput" style="display:none;" accept=".txt,.md,.py,.js,.html,.css,.json,.csv,.sh,.yaml,.yml,.toml,.ini,.cfg,.conf,.xml">
          </div>
          <div class="editor-body">
            <textarea id="docContent" class="editor-textarea" placeholder="Load or write temporary strings here..."></textarea>
            <div class="editor-ai-panel" id="editorAiPanel">
              <div class="editor-ai-header">🤖 Document AI Assistant</div>
              <div class="editor-ai-messages" id="editorAiMessages">
                <div class="editor-ai-welcome">Ask the AI to create, edit, or manage this document.</div>
              </div>
              <div class="editor-ai-input-row">
                <textarea id="editorAiInput" rows="1" placeholder="Tell the AI to edit this..."></textarea>
                <button id="editorAiSendBtn">➤</button>
              </div>
            </div>
          </div>
        </div>
      </div>

      <div class="stage-panel" id="converterPanel">
        <div class="converter-workspace">
          <h3>🔄 File Transformation Interface</h3>
          <p style="font-size: 14px;; color:var(--text-dim); margin:0 0 10px 0;">Drag and drop objects directly below. Conversions evaluate safely within isolated internal computational pipelines.</p>
          
          <div class="conversion-matrix-box">
            <div class="select-wrapper">
              <label for="targetFormatSelect">Target Configuration Format:</label>
              <select id="targetFormatSelect" class="custom-select">
                <option value="txt">Plain Text (.txt)</option>
                <option value="pdf">Portable Document Format (.pdf)</option>
                <option value="docx">Word Document (.docx)</option>
                <option value="png">Portable Network Graphics (.png)</option>
                <option value="jpeg">Joint Photographic Experts Group (.jpeg)</option>
              </select>
            </div>
            <div id="selectedFileBanner">No file staged.</div>
          </div>

          <div class="dropzone" id="converterDropzone">
            <input type="file" id="hiddenFileInput" style="display:none;" />
            <div class="dropzone-icon">📥</div>
            <div class="dropzone-text" id="dropzoneText">Drag & drop workspace file here or <span style="color:var(--violet-1); text-decoration:underline;">browse files</span></div>
            <div class="dropzone-hint">Supported matrix profiles: PDF, DOC, DOCX, TXT, PNG, JPEG</div>
          </div>

          <button class="sidebar-btn" id="executeConversionBtn" style="margin-top:10px; background:linear-gradient(135deg, var(--violet-2), var(--violet-3)); text-align:center; display:block; justify-content:center; font-weight:700;" disabled>⚡ Execute Pipeline Transformation</button>
        </div>
      </div>

      <div class="stage-panel" id="calculatorPanel">
        <div class="tool-workspace">
          <h3>🧮 Quick Calculator</h3>
          <div class="calc-display">
            <div class="calc-expr" id="calcExpr">&nbsp;</div>
            <div class="calc-value" id="calcValue">0</div>
          </div>
          <div class="calc-grid" id="calcGrid">
            <button class="calc-btn clear wide" data-act="clear">AC</button>
            <button class="calc-btn clear" data-act="back">⌫</button>
            <button class="calc-btn op" data-val="/">÷</button>

            <button class="calc-btn" data-val="7">7</button>
            <button class="calc-btn" data-val="8">8</button>
            <button class="calc-btn" data-val="9">9</button>
            <button class="calc-btn op" data-val="*">×</button>

            <button class="calc-btn" data-val="4">4</button>
            <button class="calc-btn" data-val="5">5</button>
            <button class="calc-btn" data-val="6">6</button>
            <button class="calc-btn op" data-val="-">−</button>

            <button class="calc-btn" data-val="1">1</button>
            <button class="calc-btn" data-val="2">2</button>
            <button class="calc-btn" data-val="3">3</button>
            <button class="calc-btn op" data-val="+">+</button>

            <button class="calc-btn wide" data-val="0">0</button>
            <button class="calc-btn" data-val=".">.</button>
            <button class="calc-btn eq" data-act="equals">=</button>
          </div>
        </div>
      </div>

      <div class="stage-panel" id="calendarPanel">
        <div class="calendar-workspace">
          <h3>📅 Calendar</h3>
          <div id="calNotifyBanner"></div>
          <div class="cal-header">
            <button class="cal-nav-btn" id="calPrevBtn">←</button>
            <h3 id="calMonthLabel"></h3>
            <button class="cal-nav-btn" id="calNextBtn">→</button>
          </div>
          <div class="cal-grid" id="calGrid"></div>
          <div class="cal-today-label" id="calTodayLabel"></div>

          <div class="reminder-panel">
            <div class="reminder-panel-title" id="reminderPanelTitle">Select a date to add a reminder</div>
            <div class="reminder-input-group">
              <input type="text" id="reminderInput" class="input-field" placeholder="Reminder text for selected date...">
              <button class="sidebar-btn" id="addReminderBtn" style="width:auto; margin:0;">+ Add</button>
            </div>
            <div class="reminder-item-list" id="reminderItemList"></div>
          </div>
        </div>
      </div>

    </div>
  </div>
</div>

<script>
// --- Canvas Starscape Logic ---
const canvas = document.getElementById('stars'); const ctx = canvas.getContext('2d');
let stars = [];
function resizeCanvas(){
  canvas.width = window.innerWidth; canvas.height = window.innerHeight; stars = [];
  const count = Math.floor((canvas.width * canvas.height) / 11000);
  for(let i=0;i<count;i++) stars.push({x:Math.random()*canvas.width, y:Math.random()*canvas.height, r:Math.random()*1.2+0.2, a:Math.random(), d:(Math.random()*0.01)+0.002});
}
function drawStars(){
  ctx.clearRect(0,0,canvas.width,canvas.height);
  for(const s of stars){ s.a += s.d; if(s.a>1||s.a<0) s.d*=-1; ctx.beginPath(); ctx.fillStyle=`rgba(196,181,253,${Math.abs(Math.sin(s.a))})`; ctx.arc(s.x, s.y, s.r, 0, Math.PI*2); ctx.fill(); }
  requestAnimationFrame(drawStars);
}
window.addEventListener('resize', resizeCanvas); resizeCanvas(); drawStars();

// --- Core App Management Workspace Architecture ---
let currentChatId = "";
let currentDocId = "default-doc";

const chatContainer = document.getElementById('chatContainer');
const emptyState = document.getElementById('emptyState');
const msgEl = document.getElementById('msg');
const sendBtn = document.getElementById('sendBtn');
const chatScroller = document.getElementById('chatScroller');

const scrollToBottomBtn = document.getElementById('scrollToBottomBtn');

const NEAR_BOTTOM_PX = 120;
function isNearBottom() {
  return (chatScroller.scrollHeight - chatScroller.scrollTop - chatScroller.clientHeight) < NEAR_BOTTOM_PX;
}
function scrollToBottom(force = false) {
  if (force || isNearBottom()) {
    chatScroller.scrollTop = chatScroller.scrollHeight;
  }
  updateScrollButton();
}
function updateScrollButton() {
  scrollToBottomBtn.classList.toggle('visible', !isNearBottom() && chatScroller.scrollHeight > chatScroller.clientHeight + 40);
}
chatScroller.addEventListener('scroll', updateScrollButton);
scrollToBottomBtn.addEventListener('click', () => scrollToBottom(true));
window.addEventListener('resize', updateScrollButton);

function switchView(targetPanelId, triggerElement=null) {
  document.querySelectorAll('.stage-panel').forEach(p => p.classList.remove('active'));
  document.querySelectorAll('.sidebar-btn').forEach(b => b.classList.remove('active'));
  document.querySelectorAll('.list-item-link').forEach(l => l.classList.remove('active'));
  
  document.getElementById(targetPanelId).classList.add('active');
  if(triggerElement) triggerElement.classList.add('active');
}

document.getElementById('calculatorTabLink').addEventListener('click', (e) => {
  switchView('calculatorPanel', e.currentTarget);
});

document.getElementById('calendarTabLink').addEventListener('click', (e) => {
  switchView('calendarPanel', e.currentTarget);
  fetchReminders();
});

document.getElementById('memVaultTabLink').addEventListener('click', (e) => {
  switchView('memoryPanel', e.currentTarget);
  loadMemoryWorkspace();
});
document.getElementById('docEditorTabLink').addEventListener('click', (e) => {
  switchView('editorPanel', e.currentTarget);
  loadDocument(currentDocId);
  setTimeout(() => { const el = document.getElementById('editorAiInput'); if(el) el.focus(); }, 100);
});
document.getElementById('fileConverterTabLink').addEventListener('click', (e) => {
  switchView('converterPanel', e.currentTarget);
});

async function syncWorkspaceManifest() {
  const res = await fetch('/api/init');
  const data = await res.json();
  
  const histContainer = document.getElementById('historyContainer');
  histContainer.innerHTML = "";
  data.chats.forEach(c => {
    const link = document.createElement('div');
    link.className = `list-item-link ${c.id === currentChatId ? 'active' : ''}`;
    link.textContent = "⬡ " + c.title;
    link.addEventListener('click', () => selectChatThread(c.id));
    histContainer.appendChild(link);
  });

  const docContainer = document.getElementById('savedDocsContainer');
  docContainer.innerHTML = "";
  data.documents.forEach(d => {
    const link = document.createElement('div');
    link.className = `list-item-link ${d.id === currentDocId ? 'active' : ''}`;
    link.textContent = "📝 " + d.title;
    link.addEventListener('click', () => {
      currentDocId = d.id;
      switchView('editorPanel', document.getElementById('docEditorTabLink'));
      loadDocument(d.id);
    });
    docContainer.appendChild(link);
  });
}

async function selectChatThread(chatId) {
  currentChatId = chatId;
  switchView('chatPanel');
  if(emptyState) emptyState.style.display = 'none';
  chatContainer.innerHTML = "";
  
  const res = await fetch(`/api/chat/get?chat_id=${chatId}`);
  if(res.ok) {
    const data = await res.json();
    data.steps.forEach(step => addStepNodeToStage(step));
    scrollToBottom(true);
  }
  syncWorkspaceManifest();
}

document.getElementById('newChatBtn').addEventListener('click', async () => {
  const res = await fetch('/api/chat/new', {method:'POST'});
  const data = await res.json();
  currentChatId = data.chat_id;
  if(emptyState) emptyState.style.display = 'none';
  chatContainer.innerHTML = `
    <div class="empty-state">
      <h2>Fresh Operational Thread State Enabled</h2>
      <p style="font-size: 14px;;">Long-term context memory configurations are integrated dynamically into this agent frame.</p>
    </div>`;
  syncWorkspaceManifest();
  switchView('chatPanel');
});

const QUEUE_ICONS = {pending: '☐', active: '◐', done: '✓', failed: '✕'};
let currentQueueRowEl = null;

function renderQueueInto(rowEl, queue){
  const body = rowEl.querySelector('.task-queue-body');
  body.innerHTML = "";
  queue.forEach(q => {
    const item = document.createElement('div');
    item.className = `queue-item ${q.status}`;
    const iconSpan = document.createElement('span');
    iconSpan.className = 'q-icon';
    iconSpan.textContent = QUEUE_ICONS[q.status] || '☐';
    const agentSpan = document.createElement('span');
    agentSpan.className = 'q-agent';
    agentSpan.textContent = q.agent;
    const textSpan = document.createElement('span');
    textSpan.className = 'q-text';
    textSpan.textContent = q.task;
    item.appendChild(iconSpan); item.appendChild(agentSpan); item.appendChild(textSpan);
    body.appendChild(item);
  });
}

function agentBadge(agent){
  if(!agent) return null;
  const badge = document.createElement('span');
  badge.className = `agent-badge agent-${agent}`;
  badge.textContent = agent;
  return badge;
}

function addStepNodeToStage(step) {
  const row = document.createElement('div');
  row.className = 'row assistant';
  const timeStr = step.time || '';

  if(step.type === 'user_msg') {
    currentQueueRowEl = null;
    row.className = 'row user';
    row.innerHTML = `<div class="bubble user"></div>`;
    const bubble = row.querySelector('.bubble');
    bubble.appendChild(document.createTextNode(step.text));
    if(timeStr) { const ts = document.createElement('span'); ts.className = 'msg-time-inline'; ts.textContent = timeStr; bubble.appendChild(ts); }
  } else if(step.type === 'task_queue'){
    if(currentQueueRowEl && currentQueueRowEl.isConnected){
      renderQueueInto(currentQueueRowEl, step.queue);
      scrollToBottom();
      return;
    }
    row.innerHTML = `<div class="task-queue-block"><div class="task-queue-head">📋 Task Queue</div><div class="task-queue-body"></div></div>`;
    renderQueueInto(row, step.queue);
    currentQueueRowEl = row;
  } else if(step.type === 'thought'){
    return;
  } else if(step.type === 'command'){
    row.innerHTML = `<div class="term-block"><div class="term-head"><span class="term-dots"><span></span><span></span><span></span></span></div><div class="term-body command"></div></div>`;
    const head = row.querySelector('.term-head');
    const badge = agentBadge(step.agent);
    if(badge) head.appendChild(badge);
    head.appendChild(document.createTextNode('operational syntax command execution'));
    row.querySelector('.term-body').textContent = step.text;
  } else if(step.type === 'output'){
    row.innerHTML = `<div class="term-block"><div class="term-head"><span class="term-dots"><span></span><span></span><span></span></span></div><div class="term-body output"></div></div>`;
    const head = row.querySelector('.term-head');
    const badge = agentBadge(step.agent);
    if(badge) head.appendChild(badge);
    head.appendChild(document.createTextNode('runtime captured environment buffer'));
    row.querySelector('.term-body').textContent = step.text || '(empty standard output buffer captured)';
  } else if(step.type === 'review'){
    row.innerHTML = `<div class="review-block ${step.verdict === 'PASS' ? 'pass' : 'fail'}"></div>`;
    row.querySelector('.review-block').appendChild(document.createTextNode(step.text));
  } else if(step.type === 'final'){
    row.innerHTML = `<div class="bubble final"></div>`;
    const bubble = row.querySelector('.bubble');
    const badge = agentBadge(step.agent && step.agent !== 'reviewer' ? step.agent : null);
    if(badge){ bubble.appendChild(badge); bubble.appendChild(document.createElement('br')); }
    bubble.appendChild(document.createTextNode(step.text));
    if(timeStr) { const ts = document.createElement('span'); ts.className = 'msg-time-inline'; ts.textContent = timeStr; bubble.appendChild(ts); }
  } else if(step.type === 'warning'){
    row.innerHTML = `<div style="color:var(--yellow); font-size: 13px;; padding:4px;">⚠ ${step.text}</div>`;
  } else if(step.type === 'error'){
    row.innerHTML = `<div style="color:var(--red); font-size: 13px;; padding:4px;">✕ ${step.text}</div>`;
  }
  chatContainer.appendChild(row);
  scrollToBottom();
}

async function dispatchMessage(){
  const text = msgEl.value.trim();
  if(!text) return;
  msgEl.value = ''; msgEl.style.height = 'auto'; sendBtn.disabled = true;

  const messageText = text;

  if (emptyState) emptyState.style.display = 'none';
  const now = new Date();
  const timeStr = now.toLocaleTimeString('en-US', {hour:'2-digit', minute:'2-digit'});
  addStepNodeToStage({type: 'user_msg', text: text, time: timeStr});
  scrollToBottom(true);

  let thoughtRow = null;
  let thoughtEl = null;

  function ensureThoughtEl() {
    if (!thoughtEl || !thoughtEl.isConnected) {
      thoughtRow = document.createElement('div');
      thoughtRow.className = 'row assistant';
      thoughtEl = document.createElement('div');
      thoughtEl.className = 'thought';
      thoughtEl.innerHTML = '<span class="think-dots"><span>.</span><span>.</span><span>.</span></span>';
      thoughtRow.appendChild(thoughtEl);
      chatContainer.appendChild(thoughtRow);
      scrollToBottom(true);
    }
  }

  try {
    const res = await fetch('/api/chat/stream', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({message: messageText, chat_id: currentChatId})
    });

    const reader = res.body.getReader();
    const decoder = new TextDecoder();
    let buffer = '';

    while(true) {
      const {done, value} = await reader.read();
      if(done) break;
      buffer += decoder.decode(value, {stream: true});

      const lines = buffer.split('\n');
      buffer = lines.pop();

      for(const line of lines) {
        if(!line.trim()) continue;
        try {
          const step = JSON.parse(line);

          if(step.type === 'thinking_start') {
            ensureThoughtEl();
            continue;
          }
          if(step.type === 'token') {
            ensureThoughtEl();
            if (thoughtEl) {
              const t = step.text;
              if (thoughtEl.dataset.stopped !== '1') {
                const buf = (thoughtEl.dataset.buf || '') + t;
                thoughtEl.dataset.buf = buf;
                if (!thoughtEl.dataset.cleared) {
                  thoughtEl.dataset.cleared = '1';
                  thoughtEl.textContent = '';
                }
                if (buf.includes('ACTION:') || buf.includes('ANSWER:') || buf.includes('COMMAND:') || /\n?(ACTION|ANSWER|COMMAND)$/.test(buf)) {
                  thoughtEl.dataset.stopped = '1';
                } else {
                  let fullThought = '';
                  const thIdx = buf.indexOf('THOUGHT:');
                  if (thIdx !== -1) {
                    const after = buf.slice(thIdx + 8);
                    const keys = ['\nANSWER:', '\nACTION:', '\nCOMMAND:', 'ANSWER:', 'ACTION:', 'COMMAND:'];
                    let endIdx = after.length;
                    for (const k of keys) {
                      const ki = after.indexOf(k);
                      if (ki !== -1 && ki < endIdx) endIdx = ki;
                    }
                    fullThought = after.slice(0, endIdx).trim();
                  }
                  const shown = thoughtEl.dataset.shown || '';
                  if (fullThought.length > shown.length) {
                    const newPart = fullThought.slice(shown.length);
                    thoughtEl.dataset.shown = fullThought;
                    thoughtEl.textContent += newPart.replace(/litlaw\.py/g, 'litelaw.py').replace(/litalaw\.py/g, 'litelaw.py').replace(/litalw\.py/g, 'litelaw.py');
                  }
                }
              }
              scrollToBottom();
            }
            continue;
          }
          if(step.type === 'thinking_end') {
            continue;
          }
          if(step.type === 'done') {
            currentChatId = step.chat_id;
            syncWorkspaceManifest();
            const editorPanel = document.getElementById('editorPanel');
            if (editorPanel && editorPanel.classList.contains('active')) {
              loadDocument(currentDocId);
            }
            continue;
          }
          if(step.type === 'final') {
            thoughtRow = null;
            thoughtEl = null;
            const finalRow = document.createElement('div');
            finalRow.className = 'row assistant';
            const finalBubble = document.createElement('div');
            finalBubble.className = 'bubble final';
            finalRow.appendChild(finalBubble);
            chatContainer.appendChild(finalRow);
            const fullText = (step.text || '').split('\n').filter(l => !l.startsWith('THOUGHT:')).join('\n').trim();
            let ci = 0;
            function typeChar() {
              if (ci < fullText.length) {
                finalBubble.textContent += fullText[ci];
                ci++;
                scrollToBottom();
                setTimeout(typeChar, 5);
              } else {
                const ts = document.createElement('span');
                ts.className = 'msg-time-inline';
                ts.textContent = step.time || timeStr;
                finalBubble.appendChild(ts);
                scrollToBottom();
              }
            }
            typeChar();
            continue;
          }
          thoughtRow = null;
          thoughtEl = null;
          addStepNodeToStage(step);
        } catch(e) {}
      }
    }
    thoughtRow = null;
    thoughtEl = null;
  } catch(e) {
    thoughtRow = null;
    thoughtEl = null;
    addStepNodeToStage({type:'error', text: 'Lost interface sync pipeline connection: ' + e});
  } finally {
    sendBtn.disabled = false; msgEl.focus();
  }
}

sendBtn.addEventListener('click', dispatchMessage);
msgEl.addEventListener('keydown', (e) => { if(e.key === 'Enter' && !e.shiftKey) { e.preventDefault(); dispatchMessage(); } });

// --- Long-Term Memory View Handlers ---
async function loadMemoryWorkspace() {
  const res = await fetch('/api/init');
  const data = await res.json();
  const listEl = document.getElementById('memoryCardList');
  listEl.innerHTML = "";
  
  data.memories.forEach((m, idx) => {
    const item = document.createElement('div');
    item.className = "memory-item";
    item.innerHTML = `<span>${m}</span><button class="delete-btn" data-idx="${idx}">✕ Delete</button>`;
    item.querySelector('.delete-btn').addEventListener('click', async (e) => {
      const targetIdx = e.currentTarget.getAttribute('data-idx');
      const delRes = await fetch('/api/memories/delete', {
        method:'POST', headers:{'Content-Type':'application/json'},
        body: JSON.stringify({index: parseInt(targetIdx)})
      });
      if(delRes.ok) loadMemoryWorkspace();
    });
    listEl.appendChild(item);
  });
}

document.getElementById('addMemoryBtn').addEventListener('click', async () => {
  const input = document.getElementById('memoryInput');
  const text = input.value.trim();
  if(!text) return;
  const res = await fetch('/api/memories/add', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({text: text})
  });
  if(res.ok) { input.value = ""; loadMemoryWorkspace(); }
});

// --- Document Workspace Tool Handlers ---
async function loadDocument(docId) {
  currentDocId = docId;
  const res = await fetch(`/api/documents/get?doc_id=${docId}`);
  if(res.ok) {
    const data = await res.json();
    document.getElementById('docTitle').value = data.title;
    document.getElementById('docContent').value = data.content;
  }
}

document.getElementById('saveDocBtn').addEventListener('click', async () => {
  const title = document.getElementById('docTitle').value;
  const content = document.getElementById('docContent').value;
  const res = await fetch('/api/documents/save', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({doc_id: currentDocId, title: title, content: content})
  });
  if(res.ok) { syncWorkspaceManifest(); alert("Buffer state committed successfully."); }
});

document.getElementById('newDocBtn').addEventListener('click', async () => {
  const res = await fetch('/api/documents/new', {method:'POST'});
  const data = await res.json();
  currentDocId = data.doc_id;
  loadDocument(data.doc_id);
  syncWorkspaceManifest();
});

document.getElementById('exportDocBtn').addEventListener('click', () => {
  const title = document.getElementById('docTitle').value || 'document.txt';
  const content = document.getElementById('docContent').value;
  const blob = new Blob([content], {type: 'text/plain'});
  const url = window.URL.createObjectURL(blob);
  const a = document.createElement('a');
  a.href = url;
  a.download = title;
  document.body.appendChild(a);
  a.click();
  document.body.removeChild(a);
  window.URL.revokeObjectURL(url);
});

document.getElementById('importDocBtn').addEventListener('click', () => {
  document.getElementById('importFileInput').click();
});
document.getElementById('importFileInput').addEventListener('change', async (e) => {
  const file = e.target.files[0];
  if (!file) return;
  const formData = new FormData();
  formData.append('file', file);
  const res = await fetch('/api/documents/import', {method:'POST', body: formData});
  if(res.ok) {
    const data = await res.json();
    currentDocId = data.doc_id;
    loadDocument(data.doc_id);
    syncWorkspaceManifest();
    alert(`Imported as "${data.title}"`);
  } else { alert('Import failed.'); }
  e.target.value = '';
});

// --- Editor AI Assistant ---
const editorAiInput = document.getElementById('editorAiInput');
const editorAiSendBtn = document.getElementById('editorAiSendBtn');
const editorAiMessages = document.getElementById('editorAiMessages');

async function sendEditorAiMessage() {
  const text = editorAiInput.value.trim();
  if (!text || !currentDocId) return;
  editorAiInput.value = ''; editorAiInput.style.height = 'auto';
  editorAiSendBtn.disabled = true;

  const welcome = editorAiMessages.querySelector('.editor-ai-welcome');
  if (welcome) welcome.remove();
  const userMsg = document.createElement('div');
  userMsg.className = 'editor-ai-msg user';
  userMsg.textContent = text;
  editorAiMessages.appendChild(userMsg);

  const thinking = document.createElement('div');
  thinking.className = 'editor-ai-msg assistant thought';
  thinking.textContent = '🤔 editing...';
  editorAiMessages.appendChild(thinking);
  editorAiMessages.scrollTop = editorAiMessages.scrollHeight;

  try {
    const res = await fetch('/api/chat/doc', {
      method: 'POST', headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({instruction: text, doc_id: currentDocId})
    });
    const data = await res.json();
    thinking.remove();

    if (!res.ok) {
      const errMsg = document.createElement('div');
      errMsg.className = 'editor-ai-msg assistant error';
      errMsg.textContent = '✕ ' + (data.error || 'Request failed');
      editorAiMessages.appendChild(errMsg);
    } else {
      // Update the document textarea with the new content
      document.getElementById('docContent').value = data.new_content;
      // Update the title in case it changed
      const titleInput = document.getElementById('docTitle');
      if (data.new_title) titleInput.value = data.new_title;

      const resultMsg = document.createElement('div');
      resultMsg.className = 'editor-ai-msg assistant final';
      resultMsg.textContent = data.summary || 'Document updated.';
      editorAiMessages.appendChild(resultMsg);

      syncWorkspaceManifest();
    }
  } catch (e) {
    thinking.remove();
    const errMsg = document.createElement('div');
    errMsg.className = 'editor-ai-msg assistant error';
    errMsg.textContent = '✕ Connection error: ' + e;
    editorAiMessages.appendChild(errMsg);
  } finally {
    editorAiSendBtn.disabled = false;
    editorAiInput.focus();
  }
  editorAiMessages.scrollTop = editorAiMessages.scrollHeight;
}

editorAiSendBtn.addEventListener('click', sendEditorAiMessage);
editorAiInput.addEventListener('keydown', (e) => {
  if (e.key === 'Enter' && !e.shiftKey) { e.preventDefault(); sendEditorAiMessage(); }
});


// --- File Converter Tool JavaScript Architecture ---
const dropzone = document.getElementById('converterDropzone');
const hiddenFileInput = document.getElementById('hiddenFileInput');
const selectedFileBanner = document.getElementById('selectedFileBanner');
const executeBtn = document.getElementById('executeConversionBtn');
const targetFormatSelect = document.getElementById('targetFormatSelect');
const dropzoneText = document.getElementById('dropzoneText');
let stagedFile = null;

dropzone.addEventListener('click', () => hiddenFileInput.click());

dropzone.addEventListener('dragover', (e) => {
  e.preventDefault();
  dropzone.classList.add('dragover');
});
dropzone.addEventListener('dragleave', () => {
  dropzone.classList.remove('dragover');
});
dropzone.addEventListener('drop', (e) => {
  e.preventDefault();
  dropzone.classList.remove('dragover');
  if(e.dataTransfer.files.length > 0) {
    stageUploadedFile(e.dataTransfer.files[0]);
  }
});

hiddenFileInput.addEventListener('change', (e) => {
  if(e.target.files.length > 0) {
    stageUploadedFile(e.target.files[0]);
  }
});

function stageUploadedFile(file) {
  stagedFile = file;
  selectedFileBanner.textContent = `Staged Context Frame: ${file.name} (${(file.size / 1024).toFixed(1)} KB)`;
  selectedFileBanner.style.display = 'block';
  executeBtn.disabled = false;
  dropzoneText.innerHTML = `File staged successfully. Ready for compiler transform.`;
}

executeBtn.addEventListener('click', async () => {
  if(!stagedFile) return;

  executeBtn.disabled = true;
  executeBtn.textContent = "⚙ Transforming Stream Structure...";

  const formData = new FormData();
  formData.append('file', stagedFile);
  formData.append('target_format', targetFormatSelect.value);

  try {
    const response = await fetch('/api/convert', {
      method: 'POST',
      body: formData
    });

    if(!response.ok) {
      const errorData = await response.json();
      alert(`Pipeline error: ${errorData.error || 'Unknown stream fault'}`);
    } else {
      // Catch download mapping stream object hook
      const blob = await response.blob();
      const downloadUrl = window.URL.createObjectURL(blob);
      const linkElement = document.createElement('a');
      linkElement.href = downloadUrl;
      
      const originalName = stagedFile.name.substring(0, stagedFile.name.lastIndexOf('.'));
      linkElement.download = `${originalName}.${targetFormatSelect.value}`;
      document.body.appendChild(linkElement);
      linkElement.click();
      document.body.removeChild(linkElement);
      window.URL.revokeObjectURL(downloadUrl);
    }
  } catch(e) {
    alert("Network workspace runtime error connecting to conversion stream endpoints.");
  } finally {
    executeBtn.disabled = false;
    executeBtn.textContent = "⚡ Execute Pipeline Transformation";
  }
});

// --- Calculator Tool Logic ---
let calcExprStr = "";
const calcExprEl = document.getElementById('calcExpr');
const calcValueEl = document.getElementById('calcValue');

function calcFormatDisplay(str){
  return str.replace(/\*/g, '×').replace(/\//g, '÷').replace(/-/g, '−') || '\u00A0';
}

function calcRenderPreview(){
  calcExprEl.textContent = calcFormatDisplay(calcExprStr);
  if(/[0-9]$/.test(calcExprStr)){
    try {
      const safe = calcExprStr.replace(/[^0-9+\-*/.() ]/g, '');
      const result = Function(`"use strict"; return (${safe || 0})`)();
      if(Number.isFinite(result)) calcValueEl.textContent = String(Math.round(result * 1e10) / 1e10);
    } catch(e) { }
  }
}

document.getElementById('calcGrid').addEventListener('click', (e) => {
  const btn = e.target.closest('.calc-btn');
  if(!btn) return;
  const act = btn.getAttribute('data-act');
  const val = btn.getAttribute('data-val');

  if(act === 'clear'){
    calcExprStr = "";
    calcValueEl.textContent = "0";
    calcExprEl.textContent = '\u00A0';
    return;
  }
  if(act === 'back'){
    calcExprStr = calcExprStr.slice(0, -1);
    calcRenderPreview();
    if(!calcExprStr) calcValueEl.textContent = "0";
    return;
  }
  if(act === 'equals'){
    try {
      const safe = calcExprStr.replace(/[^0-9+\-*/.() ]/g, '');
      const result = Function(`"use strict"; return (${safe || 0})`)();
      if(!Number.isFinite(result)) throw new Error('bad result');
      calcValueEl.textContent = String(Math.round(result * 1e10) / 1e10);
      calcExprEl.textContent = calcFormatDisplay(calcExprStr) + ' =';
      calcExprStr = String(result);
    } catch(e) {
      calcValueEl.textContent = "Error";
    }
    return;
  }
  if(val !== null){
    calcExprStr += val;
    calcRenderPreview();
  }
});

// --- Calendar Tool Logic ---
let calViewDate = new Date();
let calReminders = {};
let calNotifiedKey = null;
const CAL_MONTHS = ["January","February","March","April","May","June","July","August","September","October","November","December"];
const CAL_DOW = ["Su","Mo","Tu","We","Th","Fr","Sa"];

function calDateKey(d){
  const y = d.getFullYear();
  const m = String(d.getMonth() + 1).padStart(2, '0');
  const day = String(d.getDate()).padStart(2, '0');
  return `${y}-${m}-${day}`;
}
function calFormatLabel(dateKey){
  const [y, m, d] = dateKey.split('-').map(Number);
  return `${CAL_MONTHS[m - 1]} ${d}, ${y}`;
}

let calSelectedDateKey = calDateKey(new Date());

async function fetchReminders(){
  try {
    const res = await fetch('/api/reminders/get');
    if(res.ok){
      const data = await res.json();
      calReminders = data.reminders || {};
    }
  } catch(e) { }
  renderCalendar();
  renderReminderPanel();
  checkTodayReminderNotification();
}

function renderCalendar(){
  const year = calViewDate.getFullYear();
  const month = calViewDate.getMonth();
  const todayKey = calDateKey(new Date());

  document.getElementById('calMonthLabel').textContent = `${CAL_MONTHS[month]} ${year}`;
  document.getElementById('calTodayLabel').textContent = `Today: ${calFormatLabel(todayKey)}`;

  const grid = document.getElementById('calGrid');
  grid.innerHTML = "";
  CAL_DOW.forEach(d => {
    const el = document.createElement('div');
    el.className = 'cal-dow';
    el.textContent = d;
    grid.appendChild(el);
  });

  const firstDayOffset = new Date(year, month, 1).getDay();
  const daysInMonth = new Date(year, month + 1, 0).getDate();

  for(let i = 0; i < firstDayOffset; i++){
    const el = document.createElement('div');
    el.className = 'cal-day empty';
    grid.appendChild(el);
  }
  for(let day = 1; day <= daysInMonth; day++){
    const dKey = calDateKey(new Date(year, month, day));
    const el = document.createElement('div');
    const classes = ['cal-day'];
    if(dKey === todayKey) classes.push('today');
    if(dKey === calSelectedDateKey) classes.push('selected');
    if(calReminders[dKey] && calReminders[dKey].length > 0) classes.push('has-reminder');
    el.className = classes.join(' ');
    el.textContent = day;
    el.addEventListener('click', () => {
      calSelectedDateKey = dKey;
      renderCalendar();
      renderReminderPanel();
    });
    grid.appendChild(el);
  }
}

function renderReminderPanel(){
  document.getElementById('reminderPanelTitle').textContent = `Reminders for ${calFormatLabel(calSelectedDateKey)}`;
  const listEl = document.getElementById('reminderItemList');
  listEl.innerHTML = "";
  const entries = calReminders[calSelectedDateKey] || [];

  if(entries.length === 0){
    listEl.innerHTML = `<div class="reminder-item-empty">No reminders for this date yet.</div>`;
    return;
  }
  entries.forEach((text, idx) => {
    const item = document.createElement('div');
    item.className = 'reminder-item';
    item.innerHTML = `<span></span><button class="delete-btn">✕</button>`;
    item.querySelector('span').textContent = text;
    item.querySelector('.delete-btn').addEventListener('click', async () => {
      const res = await fetch('/api/reminders/delete', {
        method: 'POST', headers: {'Content-Type': 'application/json'},
        body: JSON.stringify({date: calSelectedDateKey, index: idx})
      });
      if(res.ok){
        const data = await res.json();
        calReminders = data.reminders || {};
        renderCalendar();
        renderReminderPanel();
        checkTodayReminderNotification();
      }
    });
    listEl.appendChild(item);
  });
}

document.getElementById('addReminderBtn').addEventListener('click', async () => {
  const input = document.getElementById('reminderInput');
  const text = input.value.trim();
  if(!text) return;
  const res = await fetch('/api/reminders/add', {
    method: 'POST', headers: {'Content-Type': 'application/json'},
    body: JSON.stringify({date: calSelectedDateKey, text})
  });
  if(res.ok){
    const data = await res.json();
    calReminders = data.reminders || {};
    input.value = "";
    renderCalendar();
    renderReminderPanel();
    if(calSelectedDateKey === calDateKey(new Date())) calNotifiedKey = null;
    checkTodayReminderNotification();
  }
});
document.getElementById('reminderInput').addEventListener('keydown', (e) => {
  if(e.key === 'Enter'){ e.preventDefault(); document.getElementById('addReminderBtn').click(); }
});

function checkTodayReminderNotification(){
  const todayKey = calDateKey(new Date());
  const entries = calReminders[todayKey] || [];
  const banner = document.getElementById('calNotifyBanner');

  if(entries.length === 0){
    banner.innerHTML = "";
    return;
  }
  const summary = entries.length === 1 ? entries[0] : entries.join(' • ');
  banner.innerHTML = `<div class="reminder-notify-banner"><span>🔔</span><span>Today's reminder — ${summary}</span></div>`;

  if(calNotifiedKey === todayKey) return;
  calNotifiedKey = todayKey;

  if('Notification' in window){
    if(Notification.permission === 'granted'){
      new Notification('litelaw reminder', {body: summary});
    } else if(Notification.permission !== 'denied'){
      Notification.requestPermission().then(perm => {
        if(perm === 'granted') new Notification('litelaw reminder', {body: summary});
      });
    }
  }
}

document.getElementById('calPrevBtn').addEventListener('click', () => {
  calViewDate = new Date(calViewDate.getFullYear(), calViewDate.getMonth() - 1, 1);
  renderCalendar();
});
document.getElementById('calNextBtn').addEventListener('click', () => {
  calViewDate = new Date(calViewDate.getFullYear(), calViewDate.getMonth() + 1, 1);
  renderCalendar();
});

setInterval(fetchReminders, 60000);

// Initialize Framework Workspace state configurations on startup
syncWorkspaceManifest();
fetchReminders();
switchView('chatPanel');
</script>
</body>
</html>
"""

if __name__ == "__main__":
    print("====================================================")
    print(" ⬡  litelaw agent web environment runtime dashboard")
    print(" Local Access Interface address: http://localhost:5000")
    print("====================================================")
    app.run(host="0.0.0.0", port=5000, debug=False)
